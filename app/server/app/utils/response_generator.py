import pandas as pd
from docx import Document
from fpdf import FPDF
import io
import os

def generate_response_file(file_storage, questions, answers):
    filename = file_storage.filename
    ext = os.path.splitext(filename)[1].lower()
    file_bytes = file_storage.read()

    if ext in ['.csv', '.xls', '.xlsx']:
        df = pd.read_excel(io.BytesIO(file_bytes)) if ext != '.csv' else pd.read_csv(io.BytesIO(file_bytes))
        df = df.iloc[:len(questions), :1]
        df.columns = ['Question']
        df['Answer'] = answers
        output = io.BytesIO()
        if ext == '.csv':
            df.to_csv(output, index=False)
        else:
            df.to_excel(output, index=False)
        output.seek(0)
        return output, ext

    elif ext == '.docx':
        doc = Document(io.BytesIO(file_bytes))
        new_doc = Document()
        q_index = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and q_index < len(questions) and text == questions[q_index]:
                new_doc.add_paragraph(f"Q{q_index+1}: {text}")
                new_doc.add_paragraph(f"A{q_index+1}: {answers[q_index]}")
                q_index += 1
        output = io.BytesIO()
        new_doc.save(output)
        output.seek(0)
        return output, ext

    elif ext == '.txt' or ext == '.rtf':
        output_text = ""
        for q, a in zip(questions, answers):
            output_text += f"Q: {q}\nA: {a}\n\n"
        return io.BytesIO(output_text.encode()), ext

    elif ext == '.pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for i, (q, a) in enumerate(zip(questions, answers), 1):
            pdf.multi_cell(0, 10, txt=f"Q{i}: {q}\nA{i}: {a}\n", border=0)
        output = io.BytesIO()
        pdf.output(output)
        output.seek(0)
        return output, ext

    else:
        raise ValueError("Unsupported file type for response generation.")