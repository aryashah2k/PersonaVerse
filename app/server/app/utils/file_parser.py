import io
import os
import pandas as pd
import pdfplumber
from docx import Document
from striprtf.striprtf import rtf_to_text

ALLOWED_EXTENSIONS = ['.csv', '.xls', '.xlsx', '.docx', '.txt', '.rtf', '.pdf', '.doc']

def parse_file(file_storage):
    filename = file_storage.filename
    ext = os.path.splitext(filename)[1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file format: {ext}")

    file_bytes = file_storage.read()
    questions = []

    if ext in ['.csv', '.xls', '.xlsx']:
        df = pd.read_excel(io.BytesIO(file_bytes)) if ext != '.csv' else pd.read_csv(io.BytesIO(file_bytes))
        questions = df.iloc[:, 0].dropna().astype(str).tolist()

    elif ext == '.docx' or ext == '.doc':
        document = Document(io.BytesIO(file_bytes))
        questions = [para.text.strip() for para in document.paragraphs if para.text.strip()]
        questions += [cell.text.strip() for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip() and cell.text.strip() not in questions]

    elif ext == '.rtf':
        raw_text = rtf_to_text(file_bytes.decode(errors='ignore'))
        questions = [line.strip() for line in raw_text.splitlines() if line.strip()]

    elif ext == '.txt':
        text = file_bytes.decode(errors='ignore')
        questions = [line.strip() for line in text.splitlines() if line.strip()]

    elif ext == '.pdf':
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            lines = [line.strip() for page in pdf.pages for line in page.extract_text().split('\n') if line.strip()]
            questions = []
            buffer = ""

        for line in lines:
            if buffer:
                buffer += " " + line
                questions.append(buffer.strip())
                buffer = ""
            else:
                if line.endswith(("?", "!", ".")):
                    questions.append(line)
                else:
                    buffer = line
        return questions

    if len(questions) < 1:
        raise ValueError(f"File must contain at least 1 question. Found: {len(questions)}")

    return questions